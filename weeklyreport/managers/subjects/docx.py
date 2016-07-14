""" Wrapper onto python-docx library """
from weeklyreport.interface import ReportingInterface
from weeklyreport.decorators import accepts
from weeklyreport.log import Logger
from weeklyreport.parts.list import List
from weeklyreport.configuration import Configuration

from docx import Document
from docx.shared import Inches

class Docx(object):
    """ Class for creating reports in Microsoft Word format """
    __implements__ = (ReportingInterface,)
    _client = None
    _configuration = None
    _run = None

    MAXWIDTH = 5.7

    def __init__(self):
        """ Load the driver and create title page """
        Logger().info('Initialising Microsoft Word format driver')
        self._configuration = Configuration()

    @property
    def client(self):
        """ Get the client interface """
        if self._client is None:
            template_file = (
                self._configuration.report.template if hasattr(self._configuration.report, 'template') else None
            )
            Logger().debug('Using template file ' + str(template_file))
            self._client = Document(docx=template_file)
        return self._client

    @accepts(str, style=(None, str))
    def add_paragraph(self, text, style=None):
        """
        Add a paragraph of text to the report

        @param text string
        """
        self._run = self.client.add_paragraph(str(text), style=style)

    @accepts(str, style=(None, str))
    def add_run(self, text, style=None):
        """
        Adds a run of text to the current active paragraph

        @param text string
        @param style str
        """
        if style is not None:
            setattr(self._run.add_run(text), style, True)
        else:
            self._run.add_run(text)

    @accepts(str, style=(None, str))
    def add_list(self, text, style=None):
        """
        Add a paragraph of text to the report

        @param text string
        @param style string [ListBullet,ListNumber]
        """
        paragraph = self.client.add_paragraph(str(text), style=style)
        paragraph.paragraph_format.left_indent = Inches(List.INDENT)

    @accepts(str, int)
    def add_heading(self, heading, level):
        """
        Add a heading to the report

        @param heading string
        @param level   int
        """
        self.client.add_heading(str(heading), int(level))

    @accepts(headings=(None, list), data=(None, list), style=str)
    def add_table(self, headings=None, data=None, style=''):
        """
        Add a table to the report

        @param headings list
        @param data     list   A nested lists of strings.
        @param style    string
        """
        table = self.client.add_table(rows=1, cols=len(headings), style=style)
        header_cells = table.rows[0].cells
        for i, heading in enumerate(headings):
            header_cells[i].text = str(heading)

        for row in data:
            cells = table.add_row().cells
            for i, col in enumerate(row):
                cells[i].text = str(col)


    @accepts(str, width=(int, float))
    def add_picture(self, filename, width=0):
        """
        Add a picture to the report

        @param filename string
        @param width    int
        """
        if width == 0:
            width = self.MAXWIDTH
        self.client.add_picture(filename, width=width)

    def add_page_break(self):
        """ Adds a page break to the report """
        self.client.add_page_break()

    def format_for_email(self):
        """
        Adds the document contents to a single table cell for display in email
        """
        template_file = (
            self._configuration.report.template if hasattr(self._configuration.report, 'template') else None
        )
        Logger().debug('Using template file ' + str(template_file))
        document = Document(docx=template_file)
        table = document.add_table(rows=1, cols=1)
        table.autofit = False

        cell = table.cell(0, 0)

        #pylint: disable=protected-access
        # We are extending the internals of the docx package to enable
        # this functionality. Therefore protected access is required
        cell._element[:] = self.client._body._body[:]
        cell.add_paragraph()

        self._client = document

    @accepts(str)
    def save(self, filename):
        """
        Save the file to disk

        @param filename string
        """
        self.client.save(filename)
