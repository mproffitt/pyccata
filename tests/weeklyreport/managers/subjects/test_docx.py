
from unittest import TestCase
from mock import patch, call, PropertyMock

from collections import namedtuple
from weeklyreport.configuration import Configuration
from weeklyreport.managers.report import ReportManager
from weeklyreport.managers.subjects.docx import Docx
from tests.mocks.dataproviders import DataProviders
from weeklyreport.log import Logger

from docx import Document

class TestDocx(TestCase):
    _document = None

    @patch('weeklyreport.log.Logger.log')
    def setUp(self, mock_log):
        mock_log.return_value = None
        Logger._instance = mock_log
        self._document = Document()

    def tearDown(self):
        self._document = None

    @patch('weeklyreport.configuration.Configuration._load')
    @patch('docx.document.Document.add_paragraph')
    def test_add_paragraph(self, mock_document, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_paragraph('hello world')
                mock_document.assert_called_with('hello world')

    @patch('weeklyreport.configuration.Configuration._load')
    @patch('docx.document.Document.add_heading')
    def test_add_heading(self, mock_document, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_heading('hello world', 1)
                mock_document.assert_called_with('hello world', 1)

    @patch('weeklyreport.configuration.Configuration._load')
    @patch('docx.document.Document.add_picture')
    def test_add_picture(self, mock_document, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_picture('/path/to/image.png')
                mock_document.assert_called_with('/path/to/image.png', width=5.7)

                r.add_picture('/another/path/to/image.png', width=75)
                mock_document.assert_called_with('/another/path/to/image.png', width=75)

    @patch('weeklyreport.configuration.Configuration._load')
    @patch('docx.document.Document.add_page_break')
    def test_add_page_break(self, mock_document, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_page_break()
                mock_document.assert_called_with()

    @patch('weeklyreport.configuration.Configuration._load')
    @patch('docx.document.Document.save')
    def test_save(self, mock_document, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.save('/path/to/file.docx')
                mock_document.assert_called_with('/path/to/file.docx')

    @patch('weeklyreport.configuration.Configuration._load')
    def test_maxwidth(self, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                self.assertEquals(5.7, r.maxwidth)

    @patch('weeklyreport.configuration.Configuration._load')
    @patch('docx.document.Document.add_table')
    def test_add_table(self, mock_document, mock_load):
        with patch('weeklyreport.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('weeklyreport.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                mock_document.return_value = self._document.add_table(rows=1, cols=3)
                r = ReportManager()
                headers = [
                    'Test header',
                    'Another Test Header',
                    'Value'
                ]

                data = [
                    [1, 2, 3],
                    [4, 5, 9],
                    [3, 9., 12]
                ]
                r.add_table(headings=headers, data=data, style='Test Style')
                mock_document.assert_called_with(rows=1, cols=3, style='Test Style')


