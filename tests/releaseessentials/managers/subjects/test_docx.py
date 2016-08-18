import os

from unittest import TestCase
from mock import patch, call, PropertyMock

from collections import namedtuple
from releaseessentials.configuration import Configuration
from releaseessentials.managers.report import ReportManager
from releaseessentials.managers.subjects.docx import Docx
from tests.mocks.dataproviders import DataProviders
from releaseessentials.log import Logger

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.section import CT_SectPr
from docx.oxml.text.paragraph import CT_P

class TestDocx(TestCase):
    _document = None

    @patch('argparse.ArgumentParser.parse_args')
    @patch('releaseessentials.log.Logger.log')
    def setUp(self, mock_log, mock_parser):
        mock_parser.return_value = []
        mock_log.return_value = None
        Logger._instance = mock_log
        self._document = Document()
        path = os.path.dirname(os.path.realpath(__file__ + '../../../../../'))
        self._path = os.path.join(path, os.path.join('tests', 'conf'))

    def tearDown(self):
        self._document = None

    @patch('releaseessentials.configuration.Configuration._get_locations')
    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.add_paragraph')
    def test_add_paragraph(self, mock_document, mock_load, mock_locations):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_locations.return_value = [self._path]
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_paragraph('hello world')
                mock_document.assert_called_with('hello world', style=None)

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.add_heading')
    def test_add_heading(self, mock_document, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_heading('hello world', 1)
                mock_document.assert_called_with('hello world', 1)

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.add_picture')
    def test_add_picture(self, mock_document, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_picture('/path/to/image.png')
                mock_document.assert_called_with('/path/to/image.png', width=5.7)

                r.add_picture('/another/path/to/image.png', width=75)
                mock_document.assert_called_with('/another/path/to/image.png', width=75)

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.add_page_break')
    def test_add_page_break(self, mock_document, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_page_break()
                mock_document.assert_called_with()

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.save')
    def test_save(self, mock_document, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.save('/path/to/file.docx')
                mock_document.assert_called_with('/path/to/file.docx')

    @patch('releaseessentials.configuration.Configuration._load')
    def test_maxwidth(self, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                self.assertEquals(5.7, r.maxwidth)

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.add_table')
    def test_add_table(self, mock_document, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                mock_document.return_value = self._document.add_table(rows=1, cols=3)
                r = ReportManager()
                headers = [
                    'Test header',
                    'Another Test Header',
                    'Value'
                ]

                data = [
                    [1, 2, 3],
                    [4, 5, 9],
                    [3, 9., 12]
                ]
                r.add_table(headings=headers, data=data, style='Test Style')
                mock_document.assert_called_with(rows=1, cols=3, style='Test Style')

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.document.Document.add_paragraph')
    def test_add_list(self, mock_document, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_list('This is a list item', style='ListBullet')
                mock_document.assert_called_with('This is a list item', style='ListBullet')

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.text.paragraph.Paragraph.add_run')
    def test_add_run_with_no_style(self, mock_run, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_paragraph('hello world')
                r.add_run('This is a paragraph run')
                mock_run.assert_called_with('This is a paragraph run')

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.text.paragraph.Paragraph.add_run')
    def test_add_run_with_style(self, mock_run, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_paragraph('hello world')
                r.add_run('This is a paragraph run', style='bold')
                mock_run.assert_called_with('This is a paragraph run')
                self.assertTrue(mock_run.bold)

    @patch('releaseessentials.configuration.Configuration._load')
    @patch('docx.text.paragraph.Paragraph.add_run')
    def test_format_for_email(self, mock_run, mock_load):
        with patch('releaseessentials.configuration.Configuration.reporting', new_callable=PropertyMock) as mock_reporting:
            with patch('releaseessentials.configuration.Configuration._configuration', new_callable=PropertyMock) as mock_config:
                mock_config.return_value = DataProviders._get_config_for_test_no_template()
                mock_reporting.return_value = 'docx'
                r = ReportManager()
                r.add_paragraph('hello world')
                r.add_run('This is a paragraph run', style='bold')
                mock_run.assert_called_with('This is a paragraph run')
                self.assertTrue(mock_run.bold)
                r.format_for_email()
                self.assertIsInstance(r._client._client._body._body[0], CT_Tbl)
                self.assertIsInstance(r._client._client._body._body[1], CT_SectPr)

